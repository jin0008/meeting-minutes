#!/usr/bin/env python3
"""
회의 녹음 + 자동 회의록 생성기
- 마이크로 직접 녹음하거나 기존 오디오 파일 사용
- OpenAI Whisper로 한국어 음성→텍스트 변환
- Claude 또는 Gemini API로 회의록 자동 작성 (--api 옵션으로 선택)
"""

import os
import sys
import ssl
import argparse
import datetime
import wave
from pathlib import Path

# macOS에서 SSL 인증서 오류 방지 (python.org 설치 버전 대응)
ssl._create_default_https_context = ssl._create_unverified_context

# ─────────────────────────────────────────────
# 설정
# ─────────────────────────────────────────────
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
GEMINI_API_KEY    = os.environ.get("GEMINI_API_KEY", "")
WHISPER_MODEL = "medium"   # tiny / base / small / medium / large
SAMPLE_RATE = 16000
CHANNELS = 1

# ─────────────────────────────────────────────
# 고정 참석자 정보 (파트: 이름)
# ─────────────────────────────────────────────
PARTICIPANTS = {
    "주임교수": "서경률",
    "병원장":   "김찬윤",
    "총무":     "김태임",
    "부총무":   "고재상",
    "수련":     "김용준",
    "학생":     "윤진숙",
    "연구":     "한진우",
    "수술실":   "이승규",
    "외래":     "민지상",
    "인재개발": "전익현",
    "세목회":   "이지혜",
}


def participants_info() -> str:
    """프롬프트에 삽입할 참석자 정보 문자열을 반환합니다."""
    lines = [f"  - {part} 담당: {name} 교수" for part, name in PARTICIPANTS.items()]
    return "\n".join(lines)


# ─────────────────────────────────────────────
# 공통 프롬프트 빌더
# ─────────────────────────────────────────────
def build_prompt(transcript: str, meeting_title: str, meeting_date: str) -> str:
    return f"""아래는 회의 음성을 텍스트로 변환한 내용입니다.
이 내용을 바탕으로 전문적인 회의록을 작성해주세요.

회의 제목: {meeting_title}
회의 날짜: {meeting_date}

[고정 참석자 명단]
아래 교수님들이 회의에 참여합니다. 음성에서 이름이나 파트가 언급될 경우 해당 교수님의 발언으로 귀속해주세요.
{participants_info()}

[음성 변환 원문]
{transcript}

---

다음 형식으로 회의록을 작성해주세요.
발언자가 파악되는 경우 반드시 이름(파트)을 명시하세요.
참석 여부는 원문에서 언급이 있거나 발언이 확인된 경우에만 ✅로 표시하고, 불분명하면 -로 표시하세요.

# {meeting_title}
**날짜:** {meeting_date}
**작성자:** (총무 담당: 김태임 교수)

## 1. 참석자 현황

| 파트 | 이름 | 참석 |
|------|------|------|
{chr(10).join(f"| {part} | {name} 교수 | - |" for part, name in PARTICIPANTS.items())}

## 2. 주요 안건
(회의에서 다룬 핵심 주제들을 번호 목록으로 정리)

## 3. 파트별 보고 및 논의 내용
(각 파트 담당 교수님의 보고 내용과 논의 사항을 아래 형식으로 작성.
 발언이 확인된 파트만 포함하고, 발언자 이름(파트)을 명시할 것)

예시)
### 총무 (김태임 교수)
- 보고/발언 내용 요약

### 수련 (김용준 교수)
- 보고/발언 내용 요약

## 4. 결정 사항
(회의에서 확정된 사항들을 번호 목록으로 정리. 결정 주체가 파악되면 명시)

## 5. Action Items (후속 조치)

| 담당자 | 내용 | 기한 |
|--------|------|------|
| (이름/파트) | (할 일) | (기한 또는 미정) |

## 6. 기타 / 다음 회의 예정
(기타 언급 사항, 다음 회의 일정 등)

---
*본 회의록은 AI에 의해 자동 생성되었습니다. 내용을 검토 후 활용해주세요.*
"""


# ─────────────────────────────────────────────
# 1. 녹음
# ─────────────────────────────────────────────
def record_audio(output_path: str, duration: int | None = None) -> str:
    """
    마이크로 녹음합니다.
    duration이 None이면 Enter 키를 누를 때까지 녹음합니다.
    """
    try:
        import sounddevice as sd
        import numpy as np
    except ImportError:
        print("❌ sounddevice가 설치되지 않았습니다: pip install sounddevice")
        sys.exit(1)

    print("\n🎙️  녹음을 시작합니다...")
    if duration:
        print(f"   {duration}초 동안 녹음합니다.")
    else:
        print("   녹음을 멈추려면 Enter를 누르세요.")

    frames = []

    def callback(indata, frame_count, time_info, status):
        frames.append(indata.copy())

    with sd.InputStream(samplerate=SAMPLE_RATE, channels=CHANNELS,
                        dtype="int16", callback=callback):
        if duration:
            sd.sleep(duration * 1000)
        else:
            input()  # Enter 대기

    print("✅ 녹음 완료!")

    audio_data = np.concatenate(frames, axis=0)
    with wave.open(output_path, "wb") as wf:
        wf.setnchannels(CHANNELS)
        wf.setsampwidth(2)  # int16 = 2 bytes
        wf.setframerate(SAMPLE_RATE)
        wf.writeframes(audio_data.tobytes())

    print(f"   저장됨: {output_path}")
    return output_path


# ─────────────────────────────────────────────
# 2. 음성 → 텍스트 (Whisper)
# ─────────────────────────────────────────────
def transcribe_audio(audio_path: str) -> str:
    """Whisper로 오디오를 텍스트로 변환합니다 (한국어)."""
    try:
        import whisper
    except ImportError:
        print("❌ whisper가 설치되지 않았습니다: pip install openai-whisper")
        sys.exit(1)

    print(f"\n📝 Whisper 모델 로딩 중 ({WHISPER_MODEL})... (처음 실행 시 다운로드됩니다)")
    model = whisper.load_model(WHISPER_MODEL)

    print("   음성 변환 중...")
    result = model.transcribe(audio_path, language="ko", verbose=False)
    transcript = result["text"].strip()

    print("✅ 변환 완료!")
    return transcript


# ─────────────────────────────────────────────
# 3a. 회의록 생성 - Claude API
# ─────────────────────────────────────────────
def generate_minutes_claude(transcript: str, meeting_title: str, meeting_date: str) -> str:
    """Claude API를 사용해 회의록을 생성합니다."""
    try:
        import anthropic
    except ImportError:
        print("❌ anthropic이 설치되지 않았습니다: pip install anthropic")
        sys.exit(1)

    if not ANTHROPIC_API_KEY:
        print("❌ ANTHROPIC_API_KEY 환경변수가 설정되지 않았습니다.")
        print("   export ANTHROPIC_API_KEY='your-api-key'")
        sys.exit(1)

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    print("\n🤖 Claude가 회의록을 작성 중...")

    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=4096,
        messages=[{"role": "user", "content": build_prompt(transcript, meeting_title, meeting_date)}]
    )

    minutes = message.content[0].text
    print("✅ 회의록 생성 완료! (Claude)")
    return minutes


# ─────────────────────────────────────────────
# 3b. 회의록 생성 - Gemini API
# ─────────────────────────────────────────────
def generate_minutes_gemini(transcript: str, meeting_title: str, meeting_date: str) -> str:
    """Gemini API를 사용해 회의록을 생성합니다."""
    try:
        import google.generativeai as genai
    except ImportError:
        print("❌ google-generativeai가 설치되지 않았습니다: pip install google-generativeai")
        sys.exit(1)

    if not GEMINI_API_KEY:
        print("❌ GEMINI_API_KEY 환경변수가 설정되지 않았습니다.")
        print("   export GEMINI_API_KEY='your-api-key'")
        print("   API 키 발급: https://aistudio.google.com/app/apikey")
        sys.exit(1)

    genai.configure(api_key=GEMINI_API_KEY)

    # API에서 사용 가능한 모델을 직접 조회해 자동 선택
    print("\n🔍 사용 가능한 Gemini 모델 조회 중...")
    available = [
        m.name.replace("models/", "")
        for m in genai.list_models()
        if "generateContent" in m.supported_generation_methods
        and "gemini" in m.name
        and "vision" not in m.name  # 텍스트 전용 모델 우선
    ]

    if not available:
        print("❌ 사용 가능한 Gemini 모델이 없습니다. API 키를 확인해주세요.")
        sys.exit(1)

    # 모델 선택 우선순위: pro > flash > 나머지, 숫자 높을수록 우선
    def model_priority(name: str) -> tuple:
        # 버전 숫자 추출 (예: gemini-2.0-flash → 2.0)
        import re
        ver_match = re.search(r"gemini-(\d+\.\d+|\d+)", name)
        version = float(ver_match.group(1)) if ver_match else 0.0
        is_pro = 1 if "pro" in name else 0
        is_flash = 1 if "flash" in name else 0
        is_experimental = -1 if ("exp" in name or "preview" in name) else 0
        return (version, is_pro, is_flash, is_experimental)

    available.sort(key=model_priority, reverse=True)
    selected_model = available[0]

    print(f"   선택된 모델: {selected_model}")
    print(f"   (전체 사용 가능 모델: {', '.join(available[:5])}{'...' if len(available) > 5 else ''})")
    print("\n🤖 Gemini가 회의록을 작성 중...")

    model = genai.GenerativeModel(selected_model)
    response = model.generate_content(
        build_prompt(transcript, meeting_title, meeting_date),
        generation_config=genai.types.GenerationConfig(max_output_tokens=4096)
    )

    minutes = response.text
    print("✅ 회의록 생성 완료! (Gemini)")
    return minutes


# ─────────────────────────────────────────────
# 4. 파일 저장
# ─────────────────────────────────────────────
def minutes_to_docx(minutes: str, output_path: str, meeting_title: str, meeting_date: str):
    """회의록 텍스트를 Word 문서로 변환합니다."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("❌ python-docx가 설치되지 않았습니다: pip install python-docx")
        return False

    doc = Document()

    # ── 페이지 여백 설정 ──
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # ── 기본 폰트 설정 (한국어: 맑은 고딕) ──
    doc.styles['Normal'].font.name = '맑은 고딕'
    doc.styles['Normal'].font.size = Pt(10)
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    def set_heading_style(paragraph, level: int):
        sizes   = {1: 16, 2: 13, 3: 11}
        colors  = {1: '1F4E79', 2: '2E75B6', 3: '404040'}
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(sizes.get(level, 11))
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(colors.get(level, '000000'))
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    def add_inline_text(paragraph, text: str):
        """**bold** 인라인 처리"""
        parts = text.split('**')
        for idx, part in enumerate(parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            run.bold = (idx % 2 == 1)
            run.font.name = '맑은 고딕'
            run.font.size = Pt(10)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    def add_table(doc, table_lines: list):
        """마크다운 표를 Word 표로 변환"""
        rows = [l for l in table_lines
                if not all(c in '|-: ' for c in l.replace('|', '').strip())]
        if not rows:
            return
        parsed = [[c.strip() for c in r.split('|')[1:-1]] for r in rows]
        n_cols = max(len(r) for r in parsed)
        n_rows = len(parsed)

        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.style = 'Table Grid'

        for r_idx, cells in enumerate(parsed):
            for c_idx in range(n_cols):
                cell_text = cells[c_idx] if c_idx < len(cells) else ''
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.name = '맑은 고딕'
                run.font.size = Pt(9)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                if r_idx == 0:
                    run.bold = True
                    # 헤더 행 배경색
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'D6E4F0')
                    tcPr.append(shd)
        doc.add_paragraph()

    # ── 마크다운 파싱 ──
    lines = minutes.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # H1
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading('', level=1)
            p.clear()
            run = p.add_run(line[2:].strip())
            set_heading_style(p, 1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(4)

        # H2
        elif line.startswith('## '):
            p = doc.add_heading('', level=2)
            p.clear()
            run = p.add_run(line[3:].strip())
            set_heading_style(p, 2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)

        # H3
        elif line.startswith('### '):
            p = doc.add_heading('', level=3)
            p.clear()
            run = p.add_run(line[4:].strip())
            set_heading_style(p, 3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)

        # 표
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue

        # 불릿 리스트
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline_text(p, line[2:].strip())
            p.paragraph_format.space_after = Pt(1)

        # 번호 리스트
        elif len(line) > 2 and line[0].isdigit() and line[1] in '. ':
            p = doc.add_paragraph(style='List Number')
            add_inline_text(p, line[line.index(' ')+1:].strip())
            p.paragraph_format.space_after = Pt(1)

        # 구분선
        elif line.strip().startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            border = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'AAAAAA')
            border.append(bottom)
            p._p.pPr.append(border)

        # 이탤릭 주석 (*text*)
        elif line.strip().startswith('*') and line.strip().endswith('*') and not line.strip().startswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip('*'))
            run.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.name = '맑은 고딕'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        # 빈 줄
        elif line.strip() == '':
            pass

        # 일반 텍스트
        else:
            if line.strip():
                p = doc.add_paragraph()
                add_inline_text(p, line.strip())
                p.paragraph_format.space_after = Pt(2)

        i += 1

    doc.save(output_path)
    return True


def save_output(minutes: str, transcript: str, output_dir: str, title: str):
    """회의록과 원문 텍스트를 파일로 저장합니다."""
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    safe_title    = title.replace(" ", "_").replace("/", "-")
    timestamp     = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    meeting_date  = datetime.datetime.now().strftime("%Y년 %m월 %d일")

    # Word 문서 저장
    docx_path = os.path.join(output_dir, f"{timestamp}_{safe_title}_회의록.docx")
    if minutes_to_docx(minutes, docx_path, title, meeting_date):
        print(f"\n📄 회의록 저장: {docx_path}")
    else:
        # python-docx 없을 경우 md로 폴백
        md_path = os.path.join(output_dir, f"{timestamp}_{safe_title}_회의록.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(minutes)
        print(f"\n📄 회의록 저장: {md_path}")
        docx_path = md_path

    # 원문 텍스트 저장
    transcript_path = os.path.join(output_dir, f"{timestamp}_{safe_title}_원문.txt")
    with open(transcript_path, "w", encoding="utf-8") as f:
        f.write(transcript)
    print(f"📄 원문 저장:   {transcript_path}")

    return docx_path


# ─────────────────────────────────────────────
# 메인
# ─────────────────────────────────────────────
def main():
    global WHISPER_MODEL
    parser = argparse.ArgumentParser(
        description="회의 녹음 + 자동 회의록 생성기",
        formatter_class=argparse.RawTextHelpFormatter
    )
    parser.add_argument(
        "--file", "-f",
        type=str,
        default=None,
        help="기존 오디오 파일 경로 (mp3, wav, m4a 등)\n생략하면 마이크로 직접 녹음합니다."
    )
    parser.add_argument(
        "--duration", "-d",
        type=int,
        default=None,
        help="녹음 시간(초). 생략하면 Enter를 누를 때까지 녹음합니다."
    )
    parser.add_argument(
        "--title", "-t",
        type=str,
        default="회의",
        help="회의 제목 (기본값: '회의')"
    )
    parser.add_argument(
        "--output", "-o",
        type=str,
        default="./output",
        help="결과 저장 폴더 (기본값: ./output)"
    )
    parser.add_argument(
        "--whisper-model", "-w",
        type=str,
        default=WHISPER_MODEL,
        choices=["tiny", "base", "small", "medium", "large"],
        help=f"Whisper 모델 크기 (기본값: {WHISPER_MODEL})"
    )
    parser.add_argument(
        "--api", "-a",
        type=str,
        default="claude",
        choices=["claude", "gemini"],
        help="회의록 생성에 사용할 API (기본값: claude)\n  claude - Anthropic Claude API\n  gemini - Google Gemini API"
    )
    parser.add_argument(
        "--transcript-only",
        action="store_true",
        help="음성→텍스트 변환만 수행하고 회의록은 생성하지 않습니다."
    )

    args = parser.parse_args()
    WHISPER_MODEL = args.whisper_model

    meeting_date = datetime.datetime.now().strftime("%Y년 %m월 %d일")
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M")

    api_label = "Claude" if args.api == "claude" else "Gemini"
    print("=" * 50)
    print(f"  🗒️  자동 회의록 생성기 [{api_label}]")
    print("=" * 50)

    # Step 1: 오디오 준비
    if args.file:
        audio_path = args.file
        if not os.path.exists(audio_path):
            print(f"❌ 파일을 찾을 수 없습니다: {audio_path}")
            sys.exit(1)
        print(f"📂 입력 파일: {audio_path}")
    else:
        Path(args.output).mkdir(parents=True, exist_ok=True)
        audio_path = os.path.join(args.output, f"{timestamp}_recording.wav")
        record_audio(audio_path, args.duration)

    # Step 2: 음성→텍스트
    transcript = transcribe_audio(audio_path)

    print("\n--- 변환된 텍스트 미리보기 ---")
    preview = transcript[:300] + ("..." if len(transcript) > 300 else "")
    print(preview)
    print("------------------------------")

    if args.transcript_only:
        Path(args.output).mkdir(parents=True, exist_ok=True)
        safe_title = args.title.replace(" ", "_")
        transcript_path = os.path.join(args.output, f"{timestamp}_{safe_title}_원문.txt")
        with open(transcript_path, "w", encoding="utf-8") as f:
            f.write(transcript)
        print(f"\n📄 원문 저장: {transcript_path}")
        return

    # Step 3: 회의록 생성
    if args.api == "gemini":
        minutes = generate_minutes_gemini(transcript, args.title, meeting_date)
    else:
        minutes = generate_minutes_claude(transcript, args.title, meeting_date)

    # Step 4: 저장
    save_output(minutes, transcript, args.output, args.title)

    print("\n" + "=" * 50)
    print("  ✅ 완료!")
    print("=" * 50)
    print("\n[회의록 미리보기]")
    print(minutes[:600] + ("..." if len(minutes) > 600 else ""))


if __name__ == "__main__":
    main()
